import requests
import time
import re
import random
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ==============================================================================
# 🔴 配置区域
# ==============================================================================

MY_SUB = "请输入Cookie" 
USER_ID = "请输入爬取目标的uid"

START_DATE = "请输入目标起始时间" 
END_DATE   = "请输入目标终止时间"

MAX_PAGES = 请输入最大爬取页数
FILENAME = 请输入“导出word”的路径

# ==============================================================================

# 🔥【升级1】创建一个会话，模仿浏览器保持长连接
session = requests.Session()

# 🔥【升级2】配置重试策略：如果断网或报错，自动重试 3 次
retries = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
session.mount('https://', HTTPAdapter(max_retries=retries))

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Cookie": f"SUB={MY_SUB};",
    "Referer": f"https://m.weibo.cn/u/{USER_ID}",
    "Accept": "application/json, text/plain, */*"
}
session.headers.update(headers)

def safe_request(url, params=None):
    """
    🛡️ 安全请求函数：专门处理 SSL 错误和连接中断
    """
    for i in range(3): # 最多尝试3次
        try:
            # timeout=10 表示如果10秒没反应就重试，防止卡死
            resp = session.get(url, params=params, timeout=10)
            return resp
        except Exception as e:
            print(f"    ⚠️ 网络波动 (第{i+1}次重试): {e}")
            time.sleep(5) # 报错后休息5秒再试
    return None # 3次都失败，返回空

def parse_weibo_date(time_str):
    now = datetime.now()
    if "+0800" in time_str:
        try:
            parts = time_str.split()
            month_map = {"Jan":1,"Feb":2,"Mar":3,"Apr":4,"May":5,"Jun":6,"Jul":7,"Aug":8,"Sep":9,"Oct":10,"Nov":11,"Dec":12}
            return datetime(int(parts[5]), month_map.get(parts[1], 1), int(parts[2]))
        except: pass
    if '刚刚' in time_str or '分钟' in time_str or '小时' in time_str: return now
    if '昨天' in time_str: return now - timedelta(days=1)
    if re.match(r'^\d{2}-\d{2}$', time_str): return datetime.strptime(f"{now.year}-{time_str}", "%Y-%m-%d")
    if re.match(r'^\d{4}-\d{2}-\d{2}$', time_str): return datetime.strptime(time_str, "%Y-%m-%d")
    return now

def clean_html(raw_html):
    if not raw_html: return ""
    text = re.sub(r'<br\s*/?>', '\n', raw_html)
    text = re.sub(r'<[^<]+?>', '', text)
    return text.strip()

def get_full_content(weibo_id):
    url = f"https://m.weibo.cn/statuses/extend?id={weibo_id}"
    resp = safe_request(url) # 使用安全请求
    if resp:
        try:
            data = resp.json()
            if data['ok'] == 1: return data['data']['longTextContent']
        except: pass
    return None

def get_hot_comments(weibo_id):
    url = "https://m.weibo.cn/comments/hotflow"
    params = {"id": weibo_id, "mid": weibo_id, "max_id_type": 0}
    comments = []
    resp = safe_request(url, params=params) # 使用安全请求
    if resp:
        try:
            data = resp.json()
            if data['ok'] == 1 and 'data' in data['data']:
                for item in data['data']['data']:
                    user = item['user']['screen_name']
                    text = clean_html(item['text'])
                    comments.append(f"@{user}: {text}")
        except: pass
    return comments

def save_to_word(all_data):
    doc = Document()
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.add_heading(f'微博数据 ({START_DATE} 至 {END_DATE})', 0)
    
    for i, post in enumerate(all_data, 1):
        head = doc.add_heading(f"{i}. {post['raw_time']}", level=2)
        run = head.add_run(f"  (👍{post['likes']})")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        if post['is_full']:
            run = doc.add_paragraph("[已展开全文]\n").add_run()
            run.font.color.rgb = RGBColor(0, 150, 0)
            run.bold = True
            
        doc.add_paragraph(post['text'])
        
        if post['comments']:
            doc.add_paragraph("🔥 热门评论:", style='Intense Quote')
            for c in post['comments']:
                doc.add_paragraph(c, style='List Bullet')
        doc.add_paragraph("_" * 40)
        
    doc.save(FILENAME)
    print(f"\n✅ 保存成功: {FILENAME}")

def main():
    print(f"🚀 启动稳定版爬虫 | {START_DATE} ~ {END_DATE}")
    start_dt = datetime.strptime(START_DATE, "%Y-%m-%d")
    end_dt = datetime.strptime(END_DATE, "%Y-%m-%d").replace(hour=23, minute=59)

    # 获取CID
    cid = None
    resp = safe_request("https://m.weibo.cn/api/container/getIndex", params={"type": "uid", "value": USER_ID})
    if resp:
        try:
            for tab in resp.json()['data']['tabsInfo']['tabs']:
                if tab['tab_type'] == 'weibo':
                    cid = tab['containerid']
                    break
        except: pass
    
    if not cid:
        print("❌ 无法获取 ID，请检查网络或Cookie。")
        return

    all_posts = []
    page = 1
    
    while page <= MAX_PAGES:
        print(f"📡 第 {page} 页...")
        url = "https://m.weibo.cn/api/container/getIndex"
        params = {"uid": USER_ID, "containerid": cid, "page": page}

        # 🔥 使用 safe_request 代替 requests.get
        resp = safe_request(url, params=params)
        
        if not resp:
            print("❌ 本页加载失败，尝试下一页...")
            page += 1
            continue

        try:
            data = resp.json()
            if data['ok'] == 0:
                print("🏁 到底了。")
                break
            
            cards = data['data']['cards']
            for card in cards:
                if card['card_type'] == 9:
                    mblog = card['mblog']
                    raw_time = mblog['created_at']
                    is_top = mblog.get('isTop', 0)
                    
                    post_dt = parse_weibo_date(raw_time)
                    if post_dt > end_dt: continue
                    if post_dt < start_dt:
                        if is_top: continue
                        print(f"🛑 发现旧数据 {raw_time}，停止！")
                        if all_posts: save_to_word(all_posts)
                        return

                    weibo_id = mblog['id']
                    text = mblog['text']
                    is_full = False
                    
                    if mblog.get('isLongText'):
                        print(f"   🔍 展开全文... (ID: {weibo_id})")
                        full = get_full_content(weibo_id)
                        if full:
                            text = full
                            is_full = True
                            time.sleep(1.5) # 展开全文稍微慢点
                    
                    comments = get_hot_comments(weibo_id)
                    time.sleep(0.5)

                    all_posts.append({
                        "raw_time": raw_time,
                        "likes": mblog['attitudes_count'],
                        "text": clean_html(text),
                        "is_full": is_full,
                        "comments": comments
                    })
                    print(f"   ✅ 已抓取: {clean_html(text)[:15]}...")

        except Exception as e:
            print(f"❌ 解析出错: {e}")
        
        page += 1
        # 🔥 关键：增加休息时间！防封号
        sleep_t = random.uniform(3, 6)
        print(f"💤 休息 {sleep_t:.1f} 秒...")
        time.sleep(sleep_t)

    if all_posts: save_to_word(all_posts)
    else: print("⚠️ 无数据。")

if __name__ == "__main__":
    main()
    input("按回车退出")